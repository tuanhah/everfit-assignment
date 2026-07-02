"""
CTI Report Charts — pie, bar, gauge, timeline via matplotlib → BytesIO → docx.
"""
import matplotlib
matplotlib.use("Agg")

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from io import BytesIO
from docx.shared import Inches

from cti_docx_styles import COLORS_HEX, SEVERITY_COLORS_HEX


def _save_fig_to_buffer(fig) -> BytesIO:
    """Save matplotlib figure to BytesIO buffer."""
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150, facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf


def add_finding_type_pie(doc, findings: list) -> None:
    """Pie chart: finding distribution by type."""
    if not findings:
        return

    type_counts = {}
    for f in findings:
        t = f.get("type", "unknown")
        type_counts[t] = type_counts.get(t, 0) + 1

    labels = list(type_counts.keys())
    sizes = list(type_counts.values())
    colors = [
        "#1A237E", "#0096C7", "#DC2626", "#EA580C",
        "#CA8A04", "#16A34A", "#64748B", "#8B5CF6",
    ][:len(labels)]

    fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct="%1.0f%%",
        colors=colors, startangle=90,
        pctdistance=0.75, labeldistance=1.15,
        textprops={"fontsize": 9}
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontweight("bold")
    ax.set_title("Finding Distribution by Type", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(4.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1  # CENTER


def add_severity_bar(doc, findings: list) -> None:
    """Horizontal bar chart: finding count by severity."""
    if not findings:
        return

    sev_order = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
    sev_counts = {s: 0 for s in sev_order}
    for f in findings:
        w = f.get("weight", "INFO").upper()
        if w in sev_counts:
            sev_counts[w] += 1

    labels = [s for s in sev_order if sev_counts[s] > 0]
    counts = [sev_counts[s] for s in labels]
    colors = [SEVERITY_COLORS_HEX.get(s, "#64748B") for s in labels]

    fig, ax = plt.subplots(figsize=(5.5, max(2.5, len(labels) * 0.6)), dpi=150)
    bars = ax.barh(labels, counts, color=colors, height=0.5, edgecolor="white")

    for bar, count in zip(bars, counts):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2,
                str(count), va="center", fontsize=10, fontweight="bold",
                color=COLORS_HEX["text"])

    ax.set_xlabel("Count", fontsize=10, color=COLORS_HEX["muted"])
    ax.set_title("Findings by Severity", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_risk_gauge(doc, score: int, label: str = "Overall Exposure Score") -> None:
    """Semi-circular gauge showing risk score 0-100."""
    fig, ax = plt.subplots(figsize=(4, 2.5), dpi=150)

    # Draw gauge background arcs
    theta = np.linspace(np.pi, 0, 100)
    segments = [
        (0, 25, "#16A34A"),    # green
        (25, 50, "#CA8A04"),   # yellow
        (50, 75, "#EA580C"),   # orange
        (75, 100, "#DC2626"),  # red
    ]
    for start, end, color in segments:
        t = theta[start:end]
        x_outer = 1.0 * np.cos(t)
        y_outer = 1.0 * np.sin(t)
        x_inner = 0.6 * np.cos(t)
        y_inner = 0.6 * np.sin(t)
        for i in range(len(t) - 1):
            ax.fill(
                [x_inner[i], x_outer[i], x_outer[i+1], x_inner[i+1]],
                [y_inner[i], y_outer[i], y_outer[i+1], y_inner[i+1]],
                color=color, alpha=0.3
            )

    # Draw needle
    clamped = max(0, min(100, score))
    angle = np.pi - (clamped / 100) * np.pi
    needle_x = 0.85 * np.cos(angle)
    needle_y = 0.85 * np.sin(angle)
    ax.annotate("", xy=(needle_x, needle_y), xytext=(0, 0),
                arrowprops=dict(arrowstyle="-|>", color=COLORS_HEX["text"], lw=2.5))

    # Draw filled arc up to score
    t_fill = theta[:clamped] if clamped > 0 else []
    if len(t_fill) > 1:
        for i in range(len(t_fill) - 1):
            idx = int((i / len(t_fill)) * 100)
            c = "#16A34A" if idx < 25 else "#CA8A04" if idx < 50 else "#EA580C" if idx < 75 else "#DC2626"
            x_o = 1.0 * np.cos(t_fill[i:i+2])
            y_o = 1.0 * np.sin(t_fill[i:i+2])
            x_i = 0.6 * np.cos(t_fill[i:i+2])
            y_i = 0.6 * np.sin(t_fill[i:i+2])
            ax.fill(
                [x_i[0], x_o[0], x_o[1], x_i[1]],
                [y_i[0], y_o[0], y_o[1], y_i[1]],
                color=c, alpha=0.9
            )

    # Center score text
    ax.text(0, -0.15, str(score), ha="center", va="center",
            fontsize=28, fontweight="bold", color=COLORS_HEX["text"])
    ax.text(0, -0.35, label, ha="center", va="center",
            fontsize=9, color=COLORS_HEX["muted"])

    # Scale labels
    for val, pos in [(0, np.pi), (25, 3*np.pi/4), (50, np.pi/2), (75, np.pi/4), (100, 0)]:
        x = 1.15 * np.cos(pos)
        y = 1.15 * np.sin(pos)
        ax.text(x, y, str(val), ha="center", va="center", fontsize=7,
                color=COLORS_HEX["muted"])

    ax.set_xlim(-1.4, 1.4)
    ax.set_ylim(-0.5, 1.3)
    ax.set_aspect("equal")
    ax.axis("off")

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(3.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_timeline_chart(doc, events: list) -> None:
    """Horizontal timeline chart for dated events."""
    if not events:
        return

    # Sort by date
    sorted_events = sorted(events, key=lambda e: e.get("date", ""))
    labels = [e.get("event", "")[:40] for e in sorted_events]
    dates = [e.get("date", "N/A") for e in sorted_events]

    fig, ax = plt.subplots(figsize=(6, max(2, len(labels) * 0.4)), dpi=150)

    y_pos = range(len(labels))
    ax.scatter([0.5] * len(labels), y_pos, s=80, c=COLORS_HEX["accent"],
              zorder=5, edgecolors="white", linewidths=1.5)

    for i, (date, label) in enumerate(zip(dates, labels)):
        ax.text(0.55, i, f"  {date}  —  {label}", va="center", fontsize=8,
                color=COLORS_HEX["text"])

    # Vertical line connecting dots
    if len(y_pos) > 1:
        ax.vlines(0.5, min(y_pos), max(y_pos), color=COLORS_HEX["border"],
                 linewidth=2, zorder=1)

    ax.set_xlim(0, 5)
    ax.invert_yaxis()
    ax.axis("off")
    ax.set_title("Event Timeline", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15, loc="left")

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_traffic_sources_bar(doc, traffic_sources: dict) -> None:
    """Horizontal bar chart: traffic source breakdown (direct/search/referral/social/paid)."""
    if not traffic_sources:
        return

    source_colors = {
        "direct": "#1A237E",
        "search": "#0096C7",
        "referral": "#16A34A",
        "social": "#8B5CF6",
        "paid": "#EA580C",
        "email": "#EC4899",
        "display": "#CA8A04",
    }

    labels = [k.title() for k in traffic_sources.keys()]
    values = list(traffic_sources.values())
    colors = [source_colors.get(k, "#64748B") for k in traffic_sources.keys()]

    fig, ax = plt.subplots(figsize=(5.5, max(2, len(labels) * 0.5)), dpi=150)
    bars = ax.barh(labels, values, color=colors, height=0.5, edgecolor="white")

    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height() / 2,
                f"{val}%", va="center", fontsize=10, fontweight="bold",
                color=COLORS_HEX["text"])

    ax.set_xlabel("Percentage (%)", fontsize=10, color=COLORS_HEX["muted"])
    ax.set_title("Traffic Sources", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)
    ax.set_xlim(0, max(values) * 1.3 if values else 100)
    ax.invert_yaxis()
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.tick_params(left=False)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_geographic_pie(doc, top_countries: list) -> None:
    """Pie chart: visitor geographic distribution by country."""
    if not top_countries:
        return

    labels = [c.get("country", "?") for c in top_countries]
    sizes = [c.get("share", 0) for c in top_countries]

    # Add "Other" if shares don't sum to 100
    total = sum(sizes)
    if total < 100:
        labels.append("Other")
        sizes.append(100 - total)

    geo_colors = [
        "#1A237E", "#0096C7", "#16A34A", "#EA580C",
        "#8B5CF6", "#CA8A04", "#EC4899", "#64748B",
        "#0369A1", "#92400E",
    ][:len(labels)]

    fig, ax = plt.subplots(figsize=(5, 4), dpi=150)
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct="%1.0f%%",
        colors=geo_colors, startangle=90,
        pctdistance=0.75, labeldistance=1.15,
        textprops={"fontsize": 9}
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontweight("bold")
    ax.set_title("Visitor Geography", fontsize=12, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=15)

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(4.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1
