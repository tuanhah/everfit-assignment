"""
CTI Report Diagrams — entity relationship and network topology via networkx + matplotlib.
"""
import matplotlib
matplotlib.use("Agg")

import numpy as np
# NetworkX 2.x uses np.alltrue removed in NumPy 2.0 — patch for compatibility
if not hasattr(np, "alltrue"):
    np.alltrue = np.all

import matplotlib.pyplot as plt
import networkx as nx
from io import BytesIO
from docx.shared import Inches

from cti_docx_styles import COLORS_HEX


# Entity type → color mapping
ENTITY_COLORS = {
    "person": "#1A237E",
    "username": "#0096C7",
    "email": "#8B5CF6",
    "domain": "#EA580C",
    "ip": "#64748B",
    "organization": "#16A34A",
    "phone": "#EC4899",
    "location": "#92400E",
    "asset": "#475569",
    "event": "#0369A1",
}

# Connection type → edge style
EDGE_STYLES = {
    "owns": {"style": "solid", "color": "#1A237E", "width": 2.5},
    "uses": {"style": "solid", "color": "#0096C7", "width": 1.5},
    "works_at": {"style": "dashed", "color": "#16A34A", "width": 1.5},
    "linked_to": {"style": "dotted", "color": "#64748B", "width": 1.0},
    "alias": {"style": "dashdot", "color": "#8B5CF6", "width": 1.5},
    "communicated_with": {"style": "solid", "color": "#EC4899", "width": 1.5},
}

ENTITY_ICONS = {
    "person": "[P]",
    "username": "[@]",
    "email": "[E]",
    "domain": "[D]",
    "ip": "[IP]",
    "organization": "[O]",
    "phone": "[Ph]",
    "location": "[L]",
    "asset": "[A]",
    "event": "[Ev]",
}


def _save_fig_to_buffer(fig) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=150, facecolor="white")
    buf.seek(0)
    plt.close(fig)
    return buf


def add_entity_diagram(doc, subjects: list, connections: list) -> None:
    """Render entity relationship diagram using networkx."""
    if not subjects or not connections:
        return

    G = nx.DiGraph()

    # Add nodes
    node_colors = []
    node_labels = {}
    for s in subjects:
        sid = s.get("id", s.get("label", "?"))
        stype = s.get("type", "person").lower()
        label = s.get("label", sid)
        icon = ENTITY_ICONS.get(stype, "?")
        G.add_node(sid)
        node_labels[sid] = f"{icon} {label}"
        node_colors.append(ENTITY_COLORS.get(stype, "#64748B"))

    # Add edges
    edge_colors = []
    edge_styles = []
    edge_widths = []
    edge_labels = {}
    for c in connections:
        from_id = c.get("from_id", "")
        to_id = c.get("to_id", "")
        rel = c.get("relationship", "linked_to")
        if from_id in G.nodes and to_id in G.nodes:
            G.add_edge(from_id, to_id)
            style_info = EDGE_STYLES.get(rel, EDGE_STYLES["linked_to"])
            edge_colors.append(style_info["color"])
            edge_styles.append(style_info["style"])
            edge_widths.append(style_info["width"])
            edge_labels[(from_id, to_id)] = rel

    if len(G.nodes) == 0:
        return

    # Layout
    fig, ax = plt.subplots(figsize=(8, 6), dpi=150)
    pos = nx.spring_layout(G, seed=42, k=2.0)

    # Draw edges
    for i, (u, v) in enumerate(G.edges()):
        nx.draw_networkx_edges(
            G, pos, edgelist=[(u, v)], ax=ax,
            edge_color=[edge_colors[i]] if i < len(edge_colors) else ["#64748B"],
            width=edge_widths[i] if i < len(edge_widths) else 1.0,
            style=edge_styles[i] if i < len(edge_styles) else "solid",
            arrows=True, arrowsize=15, arrowstyle="-|>",
            connectionstyle="arc3,rad=0.1"
        )

    # Draw nodes
    nx.draw_networkx_nodes(
        G, pos, ax=ax,
        node_color=node_colors,
        node_size=2000,
        alpha=0.9,
        edgecolors="white",
        linewidths=2
    )

    # Draw labels
    nx.draw_networkx_labels(
        G, pos, labels=node_labels, ax=ax,
        font_size=7, font_color="white", font_weight="bold"
    )

    # Draw edge labels
    nx.draw_networkx_edge_labels(
        G, pos, edge_labels=edge_labels, ax=ax,
        font_size=6, font_color=COLORS_HEX["muted"],
        bbox=dict(boxstyle="round,pad=0.2", facecolor="white", alpha=0.8)
    )

    ax.set_title("Entity Relationship Map", fontsize=14, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=20)
    ax.axis("off")

    # Legend
    legend_items = []
    seen_types = set()
    for s in subjects:
        stype = s.get("type", "person").lower()
        if stype not in seen_types:
            seen_types.add(stype)
            icon = ENTITY_ICONS.get(stype, "?")
            color = ENTITY_COLORS.get(stype, "#64748B")
            legend_items.append(
                plt.Line2D([0], [0], marker="o", color="w",
                          markerfacecolor=color, markersize=8,
                          label=f"{icon} {stype.title()}")
            )
    if legend_items:
        ax.legend(handles=legend_items, loc="lower left", fontsize=7,
                 framealpha=0.9, edgecolor=COLORS_HEX["border"])

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(6))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1


def add_network_topology(doc, subjects: list, connections: list) -> None:
    """Render network topology (IP/domain/infra focused)."""
    infra_types = {"domain", "ip", "organization"}
    infra_subjects = [s for s in subjects if s.get("type", "").lower() in infra_types]
    infra_ids = {s.get("id", s.get("label", "")) for s in infra_subjects}

    infra_connections = [
        c for c in connections
        if c.get("from_id") in infra_ids or c.get("to_id") in infra_ids
    ]

    if not infra_subjects:
        # Fallback: use all subjects
        infra_subjects = subjects
        infra_connections = connections

    if len(infra_subjects) < 2:
        return

    G = nx.Graph()
    node_colors = []
    node_labels = {}

    for s in infra_subjects:
        sid = s.get("id", s.get("label", "?"))
        stype = s.get("type", "domain").lower()
        label = s.get("label", sid)
        G.add_node(sid)
        node_labels[sid] = f"{label}"
        node_colors.append(ENTITY_COLORS.get(stype, "#64748B"))

    for c in infra_connections:
        from_id = c.get("from_id", "")
        to_id = c.get("to_id", "")
        if from_id in G.nodes and to_id in G.nodes:
            G.add_edge(from_id, to_id, label=c.get("relationship", ""))

    fig, ax = plt.subplots(figsize=(7, 5), dpi=150)
    pos = nx.spring_layout(G, seed=123, k=2.0)

    nx.draw_networkx_edges(G, pos, ax=ax, edge_color=COLORS_HEX["border"],
                          width=1.5, alpha=0.7)
    nx.draw_networkx_nodes(G, pos, ax=ax, node_color=node_colors,
                          node_size=1500, alpha=0.9,
                          edgecolors="white", linewidths=2)
    nx.draw_networkx_labels(G, pos, labels=node_labels, ax=ax,
                           font_size=7, font_weight="bold")

    edge_labels = nx.get_edge_attributes(G, "label")
    nx.draw_networkx_edge_labels(G, pos, edge_labels=edge_labels, ax=ax,
                                font_size=6, font_color=COLORS_HEX["muted"])

    ax.set_title("Network Topology", fontsize=14, fontweight="bold",
                 color=COLORS_HEX["primary"], pad=20)
    ax.axis("off")

    buf = _save_fig_to_buffer(fig)
    doc.add_picture(buf, width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1
