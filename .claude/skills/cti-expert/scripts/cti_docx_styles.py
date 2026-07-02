"""
CTI Report DOCX Styles — fonts, colors, heading styles, header/footer.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# --- CTI Color Palette ---
COLORS = {
    "primary": RGBColor(0x1A, 0x23, 0x7E),       # Deep navy
    "accent": RGBColor(0x00, 0x96, 0xC7),          # Cyan accent
    "critical": RGBColor(0xDC, 0x26, 0x26),        # Red
    "high": RGBColor(0xEA, 0x58, 0x0C),            # Orange
    "medium": RGBColor(0xCA, 0x8A, 0x04),          # Yellow
    "low": RGBColor(0x16, 0xA3, 0x4A),             # Green
    "info": RGBColor(0x64, 0x74, 0x8B),            # Slate
    "text": RGBColor(0x1E, 0x29, 0x3B),            # Dark text
    "muted": RGBColor(0x6B, 0x72, 0x80),           # Gray text
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "bg_light": RGBColor(0xF1, 0xF5, 0xF9),        # Light gray bg
    "border": RGBColor(0xCB, 0xD5, 0xE1),          # Light border
}

# Hex versions for matplotlib
COLORS_HEX = {
    "primary": "#1A237E",
    "accent": "#0096C7",
    "critical": "#DC2626",
    "high": "#EA580C",
    "medium": "#CA8A04",
    "low": "#16A34A",
    "info": "#64748B",
    "text": "#1E293B",
    "muted": "#6B7280",
    "bg_light": "#F1F5F9",
    "border": "#CBD5E1",
}

SEVERITY_COLORS_HEX = {
    "CRITICAL": "#DC2626",
    "HIGH": "#EA580C",
    "MEDIUM": "#CA8A04",
    "LOW": "#16A34A",
    "INFO": "#64748B",
}

SEVERITY_COLORS = {
    "CRITICAL": COLORS["critical"],
    "HIGH": COLORS["high"],
    "MEDIUM": COLORS["medium"],
    "LOW": COLORS["low"],
    "INFO": COLORS["info"],
}


def setup_styles(doc: Document) -> Document:
    """Configure document styles for CTI report."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = COLORS["text"]
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size, color in [
        ("Heading 1", 22, "primary"),
        ("Heading 2", 16, "primary"),
        ("Heading 3", 13, "accent"),
    ]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = COLORS[color]
        s.paragraph_format.space_before = Pt(18 if level == "Heading 1" else 12)
        s.paragraph_format.space_after = Pt(8)

    return doc


def add_page_number(paragraph):
    """Add page number field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def setup_header_footer(doc: Document, report_id: str, classification: str = "OPEN SOURCE"):
    """Add header with classification and footer with page numbers."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(f"CTI REPORT  |  {classification}  |  {report_id}")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]
    run.font.name = "Calibri"

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"CTI Report — {report_id}  |  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = COLORS["muted"]
    run.font.name = "Calibri"
    add_page_number(fp)

    return doc


def add_cover_page(doc: Document, data: dict) -> Document:
    """Create professional cover page with CTI Report title."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("CTI REPORT")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = "Calibri"

    # Subtitle — case label
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(data.get("case", {}).get("label", "Intelligence Summary"))
    run.font.size = Pt(18)
    run.font.color.rgb = COLORS["accent"]
    run.font.name = "Calibri"

    # Divider line
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = div_para.add_run("_" * 50)
    run.font.color.rgb = COLORS["border"]

    doc.add_paragraph()

    # Metadata table
    case = data.get("case", {})
    meta_items = [
        ("Report ID", case.get("id", "N/A")),
        ("Classification", case.get("classification", "OPEN SOURCE")),
        ("Date", case.get("date", datetime.date.today().isoformat())),
        ("Analyst", case.get("analyst", "AI-Assisted OSINT")),
        ("Subject", case.get("subject", "N/A")),
        ("Status", case.get("status", "active")),
    ]

    table = doc.add_table(rows=len(meta_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_items):
        row = table.rows[i]
        cell_l = row.cells[0]
        cell_r = row.cells[1]
        cell_l.width = Inches(2)
        cell_r.width = Inches(4)

        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.font.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = COLORS["muted"]
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_r = cell_r.paragraphs[0].add_run(value)
        run_r.font.size = Pt(11)
        run_r.font.color.rgb = COLORS["text"]

    # Remove table borders for clean look
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                "</w:tcBorders>"
            )
            tcPr.append(tcBorders)

    # Page break after cover
    doc.add_page_break()

    return doc


def add_table_of_contents(doc: Document) -> Document:
    """Add a TOC field — requires Word to refresh on open."""
    toc_heading = doc.add_paragraph("Table of Contents", style="Heading 1")

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)

    run4 = p.add_run("[Right-click and Update Field to generate TOC]")
    run4.font.color.rgb = COLORS["muted"]
    run4.font.size = Pt(9)
    run4.font.italic = True

    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

    doc.add_page_break()
    return doc


def set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)
