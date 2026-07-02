"""
CTI Report Sections — builds each report section into the DOCX document.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from cti_docx_styles import COLORS, SEVERITY_COLORS, set_cell_shading

# Map string confidence levels to numeric percentages
CONFIDENCE_MAP = {
    "VERIFIED": 95, "STRONG": 80, "MODERATE": 60,
    "WEAK": 35, "TENTATIVE": 20, "CHALLENGED": 10,
}


def _normalize_confidence(val) -> int:
    """Convert confidence to integer. Handles int, string percent, or named level."""
    if isinstance(val, (int, float)):
        return int(val)
    if isinstance(val, str):
        cleaned = val.strip().rstrip("%").upper()
        if cleaned in CONFIDENCE_MAP:
            return CONFIDENCE_MAP[cleaned]
        try:
            return int(cleaned)
        except ValueError:
            return 0
    return 0


def add_executive_summary(doc: Document, data: dict) -> None:
    """Add executive summary section."""
    doc.add_heading("Executive Summary", level=1)
    summary = data.get("executive_summary", "")
    if summary:
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(12)

    # Exposure score callout if available
    exposure = data.get("case", {}).get("exposure_score")
    if exposure is not None:
        _add_callout_box(doc, f"Overall Exposure Score: {exposure}/100",
                        _severity_from_score(exposure))


def add_subject_profile(doc: Document, data: dict) -> None:
    """Add subject profile table."""
    subjects = data.get("subjects", [])
    if not subjects:
        return

    doc.add_heading("Subject Profile", level=1)

    for s in subjects[:10]:  # Limit to first 10
        # Support both 'label' (DOCX format) and 'value'/'display_name' (engine format)
        label = s.get("label") or s.get("display_name") or s.get("value") or "Unknown"
        stype = s.get("type", "N/A")
        doc.add_heading(f"{label} ({stype})", level=2)

        conf = _normalize_confidence(s.get("confidence", 0))
        fields = [
            ("Type", stype),
            ("Confidence", f"{conf}%"),
            ("Verified", "Yes" if s.get("verified") else "No"),
            ("Aliases", ", ".join(s.get("aliases", [])) or "None"),
            ("First Seen", s.get("first_seen") or s.get("first_observed") or "N/A"),
            ("Notes", s.get("notes", "")),
        ]

        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"

        for i, (label, value) in enumerate(fields):
            row = table.rows[i]
            # Label cell
            cell_l = row.cells[0]
            run = cell_l.paragraphs[0].add_run(label)
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = COLORS["primary"]
            set_cell_shading(cell_l, "F1F5F9")

            # Value cell
            cell_r = row.cells[1]
            run = cell_r.paragraphs[0].add_run(str(value))
            run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing


def add_findings_section(doc: Document, data: dict) -> None:
    """Add findings section with styled cards."""
    findings = data.get("findings", [])
    if not findings:
        return

    doc.add_heading("Key Findings", level=1)

    for i, f in enumerate(findings, 1):
        weight = f.get("weight", f.get("severity", "INFO")).upper()
        color = SEVERITY_COLORS.get(weight, COLORS["info"])

        # Finding heading with severity badge
        ftype = f.get("type", "N/A")
        ftitle = f.get("title") or ftype.title()
        h = doc.add_heading(level=2)
        run = h.add_run(f"[{weight}] ")
        run.font.color.rgb = color
        run.font.size = Pt(13)
        run = h.add_run(f"Finding {i}: {ftitle}")
        run.font.size = Pt(13)

        conf = _normalize_confidence(f.get("confidence", 0))
        desc = f.get("description") or f.get("content") or ""
        source = f.get("source_url") or f.get("source") or "N/A"
        collected = f.get("collected_at") or f.get("recorded_at") or "N/A"

        # Finding details table
        details = [
            ("ID", f.get("id", f"FND-{i:03d}")),
            ("Type", ftype),
            ("Severity", weight),
            ("Confidence", f"{conf}%"),
            ("Description", desc),
            ("Source", source),
            ("Collected", collected),
        ]

        table = doc.add_table(rows=len(details), cols=2)
        table.style = "Table Grid"

        for j, (label, value) in enumerate(details):
            row = table.rows[j]
            cell_l = row.cells[0]
            run = cell_l.paragraphs[0].add_run(label)
            run.font.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell_l, "F1F5F9")

            cell_r = row.cells[1]
            run = cell_r.paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
            if label == "Severity":
                run.font.color.rgb = color
                run.font.bold = True

        # Tags
        tags = f.get("tags", [])
        if tags:
            p = doc.add_paragraph()
            run = p.add_run("Tags: ")
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLORS["muted"]
            run = p.add_run(", ".join(tags))
            run.font.size = Pt(9)
            run.font.color.rgb = COLORS["muted"]

        doc.add_paragraph()  # Spacing


def add_connections_section(doc: Document, data: dict) -> None:
    """Add connections table."""
    connections = data.get("connections", [])
    if not connections:
        return

    doc.add_heading("Connections & Relationships", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    headers = ["From", "Relationship", "To", "Strength"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS["white"]
        set_cell_shading(cell, "1A237E")

    # Data rows
    for c in connections:
        row = table.add_row()
        values = [
            c.get("from_id") or c.get("from", ""),
            c.get("relationship") or c.get("type", ""),
            c.get("to_id") or c.get("to", ""),
            c.get("strength") or c.get("confidence", ""),
        ]
        for i, v in enumerate(values):
            run = row.cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)


def add_source_list(doc: Document, data: dict) -> None:
    """Add source/citation table."""
    sources = data.get("sources", [])
    if not sources:
        return

    doc.add_heading("Sources", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["#", "Source", "URL", "Date Accessed"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLORS["white"]
        set_cell_shading(cell, "1A237E")

    for j, src in enumerate(sources, 1):
        row = table.add_row()
        values = [
            str(j),
            src.get("name", ""),
            src.get("url", ""),
            src.get("date", ""),
        ]
        for i, v in enumerate(values):
            run = row.cells[i].paragraphs[0].add_run(v)
            run.font.size = Pt(8)


def add_intelligence_gaps(doc: Document, data: dict) -> None:
    """Add intelligence gaps section."""
    gaps = data.get("intelligence_gaps", [])
    if not gaps:
        return

    doc.add_heading("Intelligence Gaps", level=1)
    for gap in gaps:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(gap)
        run.font.size = Pt(10)


def add_recommendations(doc: Document, data: dict) -> None:
    """Add recommended next steps."""
    recs = data.get("recommendations", [])
    if not recs:
        return

    doc.add_heading("Recommended Next Steps", level=1)
    for i, rec in enumerate(recs, 1):
        # Handle both string and object formats
        if isinstance(rec, dict):
            text = rec.get("action") or rec.get("description") or rec.get("text") or str(rec)
        else:
            text = str(rec)
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(text)
        run.font.size = Pt(10)


def add_methodology_notes(doc: Document, data: dict) -> None:
    """Add analyst caveats and methodology."""
    doc.add_heading("Analyst Caveats & Methodology", level=1)

    caveats = data.get("caveats", [
        "This report is based exclusively on publicly available information.",
        "Correlation between accounts is inferred from shared identifiers.",
        "No private data was accessed. No social engineering was employed.",
    ])

    for c in caveats:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(c)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = COLORS["muted"]

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated by CTI Expert Skill — AI-Assisted Cyber Threat Intelligence")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = COLORS["muted"]


def _add_callout_box(doc: Document, text: str, severity: str) -> None:
    """Add a colored callout box."""
    color_map = {
        "CRITICAL": "FECACA",
        "HIGH": "FED7AA",
        "MEDIUM": "FEF3C7",
        "LOW": "D1FAE5",
        "INFO": "E0E7FF",
    }
    bg = color_map.get(severity, "E0E7FF")

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    run = cell.paragraphs[0].add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)
    set_cell_shading(cell, bg)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _severity_from_score(score: int) -> str:
    if score >= 76:
        return "CRITICAL"
    elif score >= 51:
        return "HIGH"
    elif score >= 26:
        return "MEDIUM"
    else:
        return "LOW"
