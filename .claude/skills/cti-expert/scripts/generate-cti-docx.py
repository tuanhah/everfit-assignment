#!/usr/bin/env python3
"""
CTI Report DOCX Generator
Generates professional CTI reports with charts, diagrams, and styled formatting.

Usage:
    python3 generate-cti-docx.py <input.json> <output.docx>
    python3 generate-cti-docx.py <input.json>  # outputs CTI-REPORT-<id>-<date>.docx

Input: JSON file matching the CTI report data schema (see sample-cti-report-data.json)
Output: Professional .docx with cover page, TOC, charts, diagrams, styled sections.
"""
import sys
import os
import json
import subprocess
import datetime

# Auto-install dependencies
def ensure_deps():
    required = {"python-docx": "docx", "matplotlib": "matplotlib", "networkx": "networkx"}
    for pkg, mod in required.items():
        try:
            __import__(mod)
        except ImportError:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "--break-system-packages", pkg],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL
            )

ensure_deps()

# Add scripts dir to path for sibling imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from cti_docx_styles import (
    setup_styles, setup_header_footer, add_cover_page, add_table_of_contents
)
from cti_docx_charts import (
    add_finding_type_pie, add_severity_bar, add_risk_gauge, add_timeline_chart,
    add_traffic_sources_bar, add_geographic_pie
)
from cti_docx_diagrams import add_entity_diagram, add_network_topology
from cti_docx_sections import (
    add_executive_summary, add_subject_profile, add_findings_section,
    add_connections_section, add_source_list, add_intelligence_gaps,
    add_recommendations, add_methodology_notes
)


def load_data(json_path: str) -> dict:
    """Load and validate JSON input data."""
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Ensure required top-level keys exist with defaults
    data.setdefault("case", {})
    data.setdefault("subjects", [])
    data.setdefault("findings", [])
    data.setdefault("connections", [])
    data.setdefault("timeline", [])
    data.setdefault("sources", [])
    data.setdefault("intelligence_gaps", [])
    data.setdefault("recommendations", [])
    data.setdefault("executive_summary", "")

    # Default case fields
    case = data["case"]
    case.setdefault("id", "CTI-001")
    case.setdefault("label", "CTI Report")
    case.setdefault("classification", "OPEN SOURCE")
    case.setdefault("analyst", "AI-Assisted CTI")
    case.setdefault("date", datetime.date.today().isoformat())
    case.setdefault("subject", "N/A")
    case.setdefault("status", "active")

    return data


def generate_output_path(data: dict, output_path: str = None) -> str:
    """Generate output .docx path from data if not specified."""
    if output_path:
        return output_path
    case = data.get("case", {})
    case_id = case.get("id", "CTI-001")
    date = case.get("date", datetime.date.today().isoformat())
    return f"CTI-REPORT-{case_id}-{date}.docx"


def build_report(data: dict) -> Document:
    """Assemble the full CTI report document."""
    doc = Document()

    # 1. Setup styles
    setup_styles(doc)

    # 2. Cover page
    add_cover_page(doc, data)

    # 3. Header/footer (after cover to avoid header on cover)
    case = data["case"]
    setup_header_footer(doc, case.get("id", "CTI-001"), case.get("classification", "OPEN SOURCE"))

    # 4. Table of contents
    add_table_of_contents(doc)

    # 5. Executive summary
    add_executive_summary(doc, data)

    # 6. Risk gauge (if exposure score available)
    exposure = case.get("exposure_score")
    if exposure is not None:
        doc.add_heading("Risk Assessment", level=1)
        add_risk_gauge(doc, int(exposure))
        doc.add_paragraph()

    # 7. Subject profiles
    add_subject_profile(doc, data)

    # 8. Key findings
    add_findings_section(doc, data)

    # 9. Charts section
    findings = data.get("findings", [])
    if findings:
        doc.add_heading("Statistical Analysis", level=1)
        add_finding_type_pie(doc, findings)
        doc.add_paragraph()
        add_severity_bar(doc, findings)
        doc.add_paragraph()

    # 9b. Visitor stats charts (if available)
    visitor_stats = data.get("visitor_stats", {})
    if visitor_stats:
        doc.add_heading("Visitor Intelligence", level=1)
        traffic_sources = visitor_stats.get("traffic_sources", {})
        if traffic_sources:
            add_traffic_sources_bar(doc, traffic_sources)
            doc.add_paragraph()
        top_countries = visitor_stats.get("top_countries", [])
        if top_countries:
            add_geographic_pie(doc, top_countries)
            doc.add_paragraph()

    # 10. Timeline
    timeline = data.get("timeline", [])
    if timeline:
        doc.add_heading("Timeline", level=1)
        add_timeline_chart(doc, timeline)
        doc.add_paragraph()

    # 11. Entity diagram
    subjects = data.get("subjects", [])
    connections = data.get("connections", [])
    if subjects and connections:
        doc.add_heading("Entity Relationships", level=1)
        add_entity_diagram(doc, subjects, connections)
        doc.add_paragraph()

        doc.add_heading("Network Topology", level=1)
        add_network_topology(doc, subjects, connections)
        doc.add_paragraph()

    # 12. Connections table
    add_connections_section(doc, data)

    # 13. Sources
    add_source_list(doc, data)

    # 14. Intelligence gaps
    add_intelligence_gaps(doc, data)

    # 15. Recommendations
    add_recommendations(doc, data)

    # 16. Methodology notes
    add_methodology_notes(doc, data)

    return doc


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 generate-cti-docx.py <input.json> [output.docx]")
        print("       Generates a professional CTI report in DOCX format.")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(json_path):
        print(f"Error: Input file not found: {json_path}")
        sys.exit(1)

    data = load_data(json_path)
    output_path = generate_output_path(data, output_path)

    print(f"Generating CTI Report: {output_path}")
    doc = build_report(data)
    doc.save(output_path)
    print(f"CTI Report saved: {output_path}")
    print(f"  Subjects: {len(data.get('subjects', []))}")
    print(f"  Findings: {len(data.get('findings', []))}")
    print(f"  Connections: {len(data.get('connections', []))}")
    print(f"  Timeline events: {len(data.get('timeline', []))}")


if __name__ == "__main__":
    main()
